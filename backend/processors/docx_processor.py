import os
import time
import tempfile
from docx import Document
from PIL import Image

from services.ocr_factory import OCRFactory


class DOCXProcessor:
    """
    Process DOCX files.

    Features:
    1. Extract editable text.
    2. Extract text from tables.
    3. OCR embedded images.
    """

    def __init__(self, engine_name="tesseract"):
        self.engine = OCRFactory.get_engine(engine_name)

    def process(self, docx_path, language="eng"):

        start_time = time.time()

        if not os.path.exists(docx_path):
            return {
                "status": "error",
                "message": "DOCX file not found."
            }

        document = Document(docx_path)

        extracted_text = []

        image_ocr_text = []

        confidences = []

        # -------------------------------------------------
        # Extract paragraph text
        # -------------------------------------------------

        for paragraph in document.paragraphs:

            text = paragraph.text.strip()

            if text:
                extracted_text.append(text)

        # -------------------------------------------------
        # Extract table text
        # -------------------------------------------------

        for table in document.tables:

            for row in table.rows:

                row_data = []

                for cell in row.cells:

                    value = cell.text.strip()

                    if value:
                        row_data.append(value)

                if row_data:
                    extracted_text.append(" | ".join(row_data))

        # -------------------------------------------------
        # OCR Images inside DOCX
        # -------------------------------------------------

        image_number = 1

        temp_folder = "temp/docx_images"

        os.makedirs(temp_folder, exist_ok=True)

        for rel in document.part._rels.values():

            if "image" in rel.target_ref:

                image_bytes = rel.target_part.blob

                temp_image = os.path.join(
                    temp_folder,
                    f"image_{image_number}.png"
                )

                with open(temp_image, "wb") as img:

                    img.write(image_bytes)

                try:

                    Image.open(temp_image)

                    result = self.engine.extract_text(
                        temp_image,
                        language
                    )

                    if result["text"].strip():

                        image_ocr_text.append(result["text"])

                    confidences.append(result["confidence"])

                except Exception:
                    pass

                image_number += 1

        # -------------------------------------------------
        # Average Confidence
        # -------------------------------------------------

        if confidences:
            average_confidence = round(
                sum(confidences) / len(confidences),
                2
            )
        else:
            average_confidence = 100.0

        processing_time = round(
            time.time() - start_time,
            2
        )

        final_text = "\n\n".join(extracted_text)

        if image_ocr_text:

            final_text += "\n\n----- OCR From Images -----\n\n"

            final_text += "\n".join(image_ocr_text)

        return {

            "status": "success",

            "file_type": "docx",

            "pages": 1,

            "text": final_text,

            "confidence": average_confidence,

            "processing_time": processing_time

        }